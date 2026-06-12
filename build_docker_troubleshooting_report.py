from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "reports"
DOCX_PATH = OUT_DIR / "docker_test_container_troubleshooting_report.docx"
PDF_PATH = OUT_DIR / "docker_test_container_troubleshooting_report.pdf"
TODAY = datetime.now().strftime("%Y-%m-%d")


BEFORE_AFTER_ROWS = [
    (
        "테스트 실행 환경",
        "개발자 PC의 로컬 Python, 로컬 브라우저, 로컬 WebDriver 상태에 의존",
        "Python 테스트 컨테이너와 Selenium Chrome 컨테이너로 실행 환경을 분리",
        "실행 환경 차이로 생기는 재현성 문제를 줄임",
    ),
    (
        "브라우저 실행 방식",
        "conftest.py가 webdriver.Chrome(), webdriver.Edge(), webdriver.Firefox()로 로컬 브라우저 직접 실행",
        "SELENIUM_REMOTE_URL이 있으면 webdriver.Remote(...)로 Selenium 컨테이너에 접속",
        "로컬 브라우저 설치 상태와 독립적으로 Chrome headless 테스트 가능",
    ),
    (
        "테스트 대상 범위",
        "전체 tests 디렉터리 또는 수동 명령으로 파일을 골라 실행해야 함",
        "docker compose run --rm tests 명령에 quiz, ppt, deep 3개 파일만 고정",
        "사용자가 작성한 핵심 3개 테스트만 반복 실행하기 쉬움",
    ),
    (
        "의존성 설치",
        "로컬 Python 환경에 requirements.txt를 직접 설치해야 함",
        "Dockerfile에서 python:3.14-slim 기반 이미지에 requirements.txt 설치",
        "새 PC나 CI에서도 같은 Python 패키지 조합으로 실행",
    ),
    (
        "다운로드 경로",
        ".env의 Windows 다운로드 경로가 컨테이너 내부 경로와 맞지 않을 수 있음",
        "Compose에서 DOWNLOAD_DIR=/tmp/ai_helpy_chat_downloads로 덮어씀",
        "컨테이너 내부에서 유효한 다운로드 경로 사용",
    ),
    (
        "Slack/Jira 알림",
        ".env 값이 그대로 들어가면 컨테이너 검증 중 알림이 나갈 수 있음",
        "Compose에서 Slack/Jira 환경변수를 빈 값으로 덮어씀",
        "테스트 컨테이너 검증 중 불필요한 외부 알림 방지",
    ),
    (
        "Docker 실행 전제",
        "Docker Desktop이 WSL needs updating 상태면 컨테이너 실행 불가",
        "wsl --update, wsl --shutdown 후 Docker Desktop Containers 화면 진입",
        "Docker Compose 기반 테스트 실행 가능 상태 확보",
    ),
]

TROUBLESHOOTING_ROWS = [
    (
        "WSL 업데이트 필요",
        "Docker Desktop이 WSL needs updating 화면에서 멈춤",
        "관리자 PowerShell에서 wsl --update 후 wsl --shutdown 수행",
        "Docker Desktop Containers 화면 진입",
    ),
    (
        "로컬 WebDriver 의존",
        "Docker 내부 테스트 컨테이너가 로컬 Chrome을 직접 띄울 수 없음",
        "SELENIUM_REMOTE_URL 기반 Remote WebDriver 분기 추가",
        "tests/test_webdriver_factory.py 1 passed",
    ),
    (
        "Docker 권한 경고",
        "샌드박스에서 C:\\Users\\user\\.docker 접근 제한",
        "승인된 권한으로 docker compose build/run 실행",
        "이미지 빌드 및 컨테이너 테스트 실행 성공",
    ),
]

FILE_CHANGES = [
    "Dockerfile: Python 3.14 slim 기반 테스트 실행 이미지 정의",
    "docker-compose.yml: selenium 서비스와 tests 서비스 정의",
    ".dockerignore: Git, 캐시, 로그, 가상환경 등 이미지에 불필요한 파일 제외",
    "conftest.py: SELENIUM_REMOTE_URL이 있을 때 webdriver.Remote 사용",
    "tests/test_webdriver_factory.py: Remote WebDriver 분기 단위 테스트 추가",
]

COMMANDS = [
    "docker compose build tests",
    "docker compose run --rm tests",
    "docker compose down",
]

VERIFICATION_ITEMS = [
    "docker compose config --quiet: Compose 설정 문법 확인 완료",
    "pytest tests/test_webdriver_factory.py -q: 원격 WebDriver 분기 단위 테스트 1 passed",
    "docker compose build tests: 테스트 실행용 Docker 이미지 빌드 성공",
    "docker compose run --rm tests: 컨테이너에서 30개 항목 실행, 26 passed / 4 xfailed",
]

SUMMARY_ITEMS = [
    "Docker 적용 전에는 테스트 성공 여부가 로컬 브라우저와 로컬 Python 상태에 크게 묶여 있었다.",
    "Docker 적용 후에는 테스트 실행 컨테이너와 브라우저 컨테이너가 같은 Compose 네트워크에서 동작한다.",
    "현재 컨테이너화 범위는 사용자가 작성한 quiz, ppt, deep 3개 테스트 파일로 제한했다.",
    "검증 결과는 26 passed, 4 xfailed이며, xfailed는 기존 테스트에서 의도된 xfail 항목으로 분류된다.",
]

NEXT_STEPS = [
    "Selenium 이미지는 latest 대신 고정 태그로 전환하면 장기 재현성이 더 좋아진다.",
    "CI에서도 docker compose run --rm tests를 사용하면 로컬과 CI 실행 방식을 맞출 수 있다.",
    "Edge/Firefox까지 확장하려면 Selenium 공식 Edge/Firefox 이미지를 서비스로 추가한다.",
]


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def margin_cell(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def width_table(table):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")


def setup_docx(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def add_docx_title(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Docker 테스트 컨테이너화 전후 비교 트러블슈팅")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    meta = doc.add_paragraph()
    meta.add_run(f"대상: AI_Helpy_chat / 작성일: {TODAY} / 범위: quiz, ppt, deep 테스트")


def add_docx_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_docx_table(doc, title, headers, rows):
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    width_table(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        shade_cell(cell, "F2F4F7")

    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margin_cell(cell)


def build_docx():
    doc = Document()
    setup_docx(doc)
    add_docx_title(doc)
    doc.add_heading("핵심 요약", level=1)
    add_docx_bullets(doc, SUMMARY_ITEMS)
    add_docx_table(
        doc,
        "사용 전 / 사용 후 비교",
        ["항목", "Docker 사용 전", "Docker 사용 후", "개선 효과"],
        BEFORE_AFTER_ROWS,
    )
    add_docx_table(
        doc,
        "트러블슈팅 포인트",
        ["문제", "사용 전 증상", "조치", "사용 후 확인"],
        TROUBLESHOOTING_ROWS,
    )
    doc.add_heading("변경 파일", level=1)
    add_docx_bullets(doc, FILE_CHANGES)
    doc.add_heading("실행 명령", level=1)
    add_docx_bullets(doc, COMMANDS)
    doc.add_heading("검증 결과", level=1)
    add_docx_bullets(doc, VERIFICATION_ITEMS)
    doc.add_heading("후속 권장사항", level=1)
    add_docx_bullets(doc, NEXT_STEPS)
    doc.save(DOCX_PATH)


def pdf_styles():
    pdfmetrics.registerFont(TTFont("Malgun", "C:/Windows/Fonts/malgun.ttf"))
    pdfmetrics.registerFont(TTFont("Malgun-Bold", "C:/Windows/Fonts/malgunbd.ttf"))
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "TitleKo",
            parent=base["Title"],
            fontName="Malgun-Bold",
            fontSize=19,
            leading=25,
            textColor=colors.HexColor("#0B2545"),
            alignment=TA_LEFT,
            spaceAfter=8,
        ),
        "meta": ParagraphStyle(
            "MetaKo",
            parent=base["Normal"],
            fontName="Malgun",
            fontSize=9,
            leading=13,
            textColor=colors.HexColor("#555555"),
            spaceAfter=12,
        ),
        "h1": ParagraphStyle(
            "H1Ko",
            parent=base["Heading1"],
            fontName="Malgun-Bold",
            fontSize=13.5,
            leading=17,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=10,
            spaceAfter=6,
        ),
        "body": ParagraphStyle(
            "BodyKo",
            parent=base["BodyText"],
            fontName="Malgun",
            fontSize=9.2,
            leading=12.8,
            spaceAfter=5,
        ),
        "small": ParagraphStyle(
            "SmallKo",
            parent=base["BodyText"],
            fontName="Malgun",
            fontSize=7.5,
            leading=10.5,
            spaceAfter=2,
        ),
        "code": ParagraphStyle(
            "CodeKo",
            parent=base["Code"],
            fontName="Courier",
            fontSize=8.3,
            leading=11,
            backColor=colors.HexColor("#F4F6F9"),
            borderColor=colors.HexColor("#DADCE0"),
            borderWidth=0.3,
            borderPadding=5,
            spaceAfter=5,
        ),
    }


def paragraph(text, style):
    escaped = (
        str(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )
    return Paragraph(escaped, style)


def pdf_bullets(items, styles):
    return ListFlowable(
        [ListItem(paragraph(item, styles["body"]), leftIndent=12) for item in items],
        bulletType="bullet",
        leftIndent=16,
    )


def pdf_table(headers, rows, widths, styles):
    data = [[paragraph(header, styles["small"]) for header in headers]]
    for row in rows:
        data.append([paragraph(cell, styles["small"]) for cell in row])
    table = Table(data, colWidths=widths, repeatRows=1, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#DADCE0")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("FONTNAME", (0, 0), (-1, 0), "Malgun-Bold"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def build_pdf():
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=LETTER,
        rightMargin=0.58 * inch,
        leftMargin=0.58 * inch,
        topMargin=0.62 * inch,
        bottomMargin=0.62 * inch,
        title="Docker 테스트 컨테이너화 전후 비교 트러블슈팅",
    )
    story = [
        paragraph("Docker 테스트 컨테이너화 전후 비교 트러블슈팅", styles["title"]),
        paragraph(f"대상: AI_Helpy_chat / 작성일: {TODAY} / 범위: quiz, ppt, deep 테스트", styles["meta"]),
        paragraph("핵심 요약", styles["h1"]),
        pdf_bullets(SUMMARY_ITEMS, styles),
        Spacer(1, 8),
        paragraph("사용 전 / 사용 후 비교", styles["h1"]),
        pdf_table(
            ["항목", "Docker 사용 전", "Docker 사용 후", "개선 효과"],
            BEFORE_AFTER_ROWS,
            [1.0 * inch, 1.8 * inch, 1.8 * inch, 1.55 * inch],
            styles,
        ),
        PageBreak(),
        paragraph("트러블슈팅 포인트", styles["h1"]),
        pdf_table(
            ["문제", "사용 전 증상", "조치", "사용 후 확인"],
            TROUBLESHOOTING_ROWS,
            [1.05 * inch, 1.8 * inch, 1.8 * inch, 1.5 * inch],
            styles,
        ),
        Spacer(1, 8),
        paragraph("변경 파일", styles["h1"]),
        pdf_bullets(FILE_CHANGES, styles),
        paragraph("실행 명령", styles["h1"]),
    ]
    for command in COMMANDS:
        story.append(paragraph(command, styles["code"]))
    story.extend(
        [
            paragraph("검증 결과", styles["h1"]),
            pdf_bullets(VERIFICATION_ITEMS, styles),
            paragraph("후속 권장사항", styles["h1"]),
            pdf_bullets(NEXT_STEPS, styles),
        ]
    )
    doc.build(story)


def build():
    OUT_DIR.mkdir(exist_ok=True)
    build_docx()
    build_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)


if __name__ == "__main__":
    build()
